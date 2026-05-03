#!/usr/bin/env python3
"""生成变式训练卷答案Word文档"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def sf(run):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def title(doc, text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p

def heading(doc, text, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p

def answer_line(doc, num, ans, size=11):
    p = doc.add_paragraph()
    r = p.add_run(f'{num}．{ans}')
    r.font.size = Pt(size)
    sf(r)
    return p

# ===== 标题 =====
title(doc, '变式训练卷答案', size=16)
doc.add_paragraph()

# ===== 一、单项选择题 =====
heading(doc, '一、单项选择题（共40小题）', size=12)
doc.add_paragraph()

choices_answers = [
    ('1', 'D'),
    ('2', 'A'),
    ('3', 'A'),
    ('4', 'A'),
    ('5', 'D'),
    ('6', 'D'),
    ('7', 'C'),
    ('8', 'D'),
    ('9', 'A'),
    ('10', 'A'),
    ('11', 'C'),
    ('12', 'B'),
    ('13', 'C'),
    ('14', 'D'),
    ('15', 'C'),
    ('16', 'A'),
    ('17', 'B'),
    ('18', 'C'),
    ('19', 'D'),
    ('20', 'C'),
    ('21', 'D'),
    ('22', 'C'),
    ('23', 'B'),
    ('24', 'C'),
    ('25', 'B'),
    ('26', 'D'),
    ('27', 'B'),
    ('28', 'A'),
    ('29', 'A'),
    ('30', 'B'),
    ('31', 'C'),
    ('32', 'D'),
    ('33', 'C'),
    ('34', 'A'),
    ('35', 'C'),
    ('36', 'A'),
    ('37', 'C'),
    ('38', 'B'),
    ('39', 'C'),
    ('40', 'C'),
]

# 分两列显示
table = doc.add_table(rows=20, cols=2)
table.style = 'Table Grid'
for idx, (num, ans) in enumerate(choices_answers):
    row = idx // 2
    col = idx % 2
    cell = table.rows[row].cells[col]
    p = cell.paragraphs[0]
    r = p.add_run(f'{num}．{ans}')
    r.font.size = Pt(11)
    sf(r)

doc.add_paragraph()

# ===== 二、解答题 =====
heading(doc, '二、解答题（共6小题）', size=12)
doc.add_paragraph()

answers_short = [
    ('41', [
        '（1）可以；种类和数量',
        '（2）协助扩散；顺浓度梯度运输（或：运输速率较快，或：不需要消耗能量）',
        '（3）内质网；信息交流（或：细胞间信息交流）',
    ]),
    ('42', [
        '（1）干旱处理时间',
        '（2）类囊体薄膜（或：基粒）；无水乙醇（或：丙酮）；CO₂的固定',
        '（3）品种B；在干旱处理10天时，品种B的绿叶率仍保持较高水平',
    ]),
    ('43', [
        '（1）专一',
        '（2）淀粉水解产物的生成量（或：麦芽糖的生成量）；温度、pH、底物浓度、酶浓度等',
        '（3）血糖浓度',
    ]),
    ('44', [
        '（1）卵母细胞（或：次级卵母细胞）；细胞中无同源染色体，且染色体着丝点已分裂',
        '（2）2；卵细胞或（第二）极体',
        '（3）基因重组',
    ]),
    ('45', [
        '（1）易于区分的相对性状 / 繁殖速度快 / 后代数目多 / 子代繁殖能力强（任答一点）',
        '（2）乙；自交后代发生性状分离的为杂合子（F₁中黄色:白色=3:1说明黄色为显性）',
        '（3）AABB×aabb或AAbb×aaBB；黄色圆粒:黄色皱粒:白色圆粒:白色皱粒=9:3:3:1',
    ]),
    ('46', [
        '（1）DNA复制；半保留复制',
        '（2）翻译；mRNA',
        '（3）抑制细菌RNA聚合酶的活性，从而抑制RNA的合成',
    ]),
]

for num, lines in answers_short:
    p = doc.add_paragraph()
    r = p.add_run(f'{num}．')
    r.font.size = Pt(11)
    r.font.bold = True
    sf(r)
    for line in lines:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(11)
        sf(r)
        p.paragraph_format.left_indent = Cm(0.74)

# 保存
out = '/home/hochoy/.openclaw/workspace-biology/变式训练卷-答案.docx'
doc.save(out)
print(f'答案已保存: {out}')

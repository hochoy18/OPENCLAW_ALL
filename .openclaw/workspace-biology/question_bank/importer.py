"""
importer.py - Excel/Word 导入器模块
"""

import os
import re
from typing import List, Dict, Optional
import openpyxl
from docx import Document


def import_from_excel(file_path: str, sheet_name: int = 0) -> List[Dict]:
    """
    从 Excel 文件导入题目
    
    参数:
        file_path: Excel 文件路径
        sheet_name: 工作表名称或索引，默认第一个工作表
    
    返回:
        题目列表，每道题为包含以下字段的字典:
        - title: 题干
        - answer: 答案
        - difficulty: 难度 (1-5)
        - source: 来源
        - knowledge_point: 知识点
        - keywords: 关键词
        - question_type: 题型
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    wb = openpyxl.load_workbook(file_path)
    
    # 获取工作表
    if isinstance(sheet_name, int):
        ws = wb.worksheets[sheet_name]
    else:
        ws = wb[sheet_name]
    
    questions = []
    
    # 从第二行开始读取（假设第一行是表头）
    for row in ws.iter_rows(min_row=2, values_only=True):
        # 跳过空行
        if not row[0] or str(row[0]).strip() == "":
            continue
        
        # 解析各列数据
        title = str(row[0]).strip() if row[0] else ""
        answer = str(row[1]).strip() if len(row) > 1 and row[1] else ""
        difficulty = _parse_difficulty(row[2] if len(row) > 2 else None)
        source = str(row[3]).strip() if len(row) > 3 and row[3] else ""
        knowledge_point = str(row[4]).strip() if len(row) > 4 and row[4] else ""
        keywords = str(row[5]).strip() if len(row) > 5 and row[5] else ""
        question_type = _parse_question_type(row[6] if len(row) > 6 else None)
        
        if title:
            questions.append({
                "title": title,
                "answer": answer,
                "difficulty": difficulty,
                "source": source,
                "knowledge_point": knowledge_point,
                "keywords": keywords,
                "question_type": question_type
            })
    
    wb.close()
    return questions


def import_from_docx(file_path: str) -> List[Dict]:
    """
    从 Word 文件导入题目
    
    参数:
        file_path: Word 文件路径
    
    返回:
        题目列表
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    doc = Document(file_path)
    questions = []
    
    current_question = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 检测是否为新题目（常见格式：以序号开头，如"1."、"（1）"、"【题目1】"等）
        if _is_question_start(text):
            # 保存之前的题目
            if current_question and current_question.get("title"):
                questions.append(current_question)
            
            # 开始新题目
            current_question = {
                "title": _clean_question_text(text),
                "answer": "",
                "difficulty": 3,
                "source": "",
                "knowledge_point": "",
                "keywords": "",
                "question_type": _detect_question_type(text)
            }
        
        # 检测是否为答案
        elif current_question and _is_answer_text(text):
            current_question["answer"] = _clean_answer_text(text)
        
        # 如果当前有题目且文本不是太长，可能是题干的补充
        elif current_question and current_question.get("title") and len(text) < 500:
            current_question["title"] += "\n" + text
    
    # 保存最后一个题目
    if current_question and current_question.get("title"):
        questions.append(current_question)
    
    # 处理表格中的题目（如果有）
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if cells and cells[0]:
                # 假设第一列是题干
                questions.append({
                    "title": cells[0],
                    "answer": cells[1] if len(cells) > 1 else "",
                    "difficulty": 3,
                    "source": "",
                    "knowledge_point": "",
                    "keywords": "",
                    "question_type": "选择题"
                })
    
    return questions


def _parse_difficulty(value) -> int:
    """解析难度值"""
    if value is None:
        return 3
    
    val_str = str(value).strip()
    
    # 数字直接返回
    if val_str.isdigit():
        return max(1, min(5, int(val_str)))
    
    # 中文数字
    cn_map = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "难": 3, "中": 3, "易": 1, "难": 5}
    for cn, num in cn_map.items():
        if cn in val_str:
            return num
    
    # 星级格式
    star_match = re.search(r"★+", val_str)
    if star_match:
        return len(star_match.group())
    
    return 3


def _parse_question_type(value) -> str:
    """解析题型"""
    if value is None:
        return "选择题"
    
    val_str = str(value).strip()
    
    if "选择" in val_str:
        return "选择题"
    elif "填空" in val_str:
        return "填空题"
    elif "解答" in val_str or "问答" in val_str:
        return "解答题"
    elif "判断" in val_str:
        return "判断题"
    elif "实验" in val_str:
        return "实验题"
    
    return val_str if val_str else "选择题"


def _is_question_start(text: str) -> bool:
    """判断文本是否是新题目的开始"""
    # 常见序号格式
    patterns = [
        r"^【(.+?)】",           # 【题目1】
        r"^\d+[\.、\．]",        # 1.  1、  1．
        r"^\([\d一二三四五六七八九十]+\)",  # (1) (一)
        r"^[\d一二三四五六七八九十]+、",   # 1、  一、
        r"^第[\d一二三四五六七八九十]+[题题]",  # 第1题
        r"^[A-Z][\.\、]",         # A.  A、
        r"^【\d+】",              # 【1】
    ]
    
    for pattern in patterns:
        if re.match(pattern, text):
            return True
    
    # 如果文本很短且以特定关键词开头
    if len(text) < 200 and any(kw in text for kw in ["以下", "关于", "细胞", "DNA", "蛋白质", "光合作用"]):
        return True
    
    return False


def _is_answer_text(text: str) -> bool:
    """判断文本是否为答案"""
    answer_keywords = ["答案", "答：", "答曰", "正确答案", "参考解答"]
    return any(kw in text for kw in answer_keywords)


def _clean_question_text(text: str) -> str:
    """清理题目文本"""
    # 去除序号部分
    patterns = [
        r"^【[^】]+】\s*",
        r"^\d+[\.、\．]\s*",
        r"^\([\d一二三四五六七八九十]+\)\s*",
        r"^[\d一二三四五六七八九十]+、\s*",
        r"^第[\d一二三四五六七八九十]+[题题]\s*",
    ]
    
    result = text
    for pattern in patterns:
        result = re.sub(pattern, "", result)
    
    return result.strip()


def _clean_answer_text(text: str) -> str:
    """清理答案文本"""
    # 去除"答案："等前缀
    result = re.sub(r"^(答案|答：|答曰|正确答案|参考解答)\s*", "", text)
    return result.strip()


def _detect_question_type(text: str) -> str:
    """根据文本内容推测题型"""
    if any(kw in text for kw in ["以下哪项", "正确的是", "错误的是", "属于", "包括"]):
        return "选择题"
    elif any(kw in text for kw in ["填空", "填入", "补充"]):
        return "填空题"
    elif any(kw in text for kw in ["实验", "探究", "设计"]):
        return "实验题"
    else:
        return "解答题"


def import_from_json(file_path: str) -> List[Dict]:
    """从 JSON 文件导入题目"""
    import json
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    if isinstance(data, list):
        return data
    elif isinstance(data, dict) and "questions" in data:
        return data["questions"]
    
    return []


def import_questions_auto(file_path: str) -> List[Dict]:
    """
    自动识别文件类型并导入题目
    
    支持的文件类型:
    - .xlsx, .xls: Excel 文件
    - .docx: Word 文件
    - .json: JSON 文件
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext in [".xlsx", ".xls"]:
        return import_from_excel(file_path)
    elif ext == ".docx":
        return import_from_docx(file_path)
    elif ext == ".json":
        return import_from_json(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")

"""
exporter.py - Word 试卷导出器模块
"""

import os
from typing import List, Dict
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def export_to_word(questions: List[Dict], output_path: str, title: str = "高中生物试卷") -> str:
    """
    将选中的题目导出为 Word 文档
    
    参数:
        questions: 题目列表
        output_path: 输出文件路径
        title: 试卷标题
    
    返回:
        生成的 Word 文件路径
    """
    doc = Document()
    
    # 设置默认字体（支持中文）
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')
    style.font.size = Pt(12)
    
    # 添加标题
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加试卷信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_run = info_para.add_run("姓名：___________    学号：___________    得分：___________")
    info_run.font.size = Pt(11)
    
    doc.add_paragraph()  # 空行
    
    # 按题型分组
    questions_by_type = {}
    for q in questions:
        qtype = q.get("question_type", "选择题")
        if qtype not in questions_by_type:
            questions_by_type[qtype] = []
        questions_by_type[qtype].append(q)
    
    # 题目序号
    q_index = 1
    
    # 写入每种题型
    for qtype, type_questions in questions_by_type.items():
        # 题型标题
        type_heading = doc.add_heading(qtype, level=1)
        
        for q in type_questions:
            # 题干
            q_para = doc.add_paragraph()
            q_text = f"{q_index}. "
            
            # 添加难度标记
            difficulty = q.get("difficulty", 3)
            if difficulty:
                stars = "★" * difficulty + "☆" * (5 - difficulty)
                q_text += f"[{stars}] "
            
            q_text += q.get("title", "")
            q_run = q_para.add_run(q_text)
            q_run.font.size = Pt(12)
            
            # 如果是选择题，添加选项格式
            if qtype == "选择题" and "\n" in q.get("title", ""):
                # 分割选项
                lines = q.get("title", "").split("\n")
                for line in lines[1:]:  # 跳过第一行（题干）
                    if line.strip():
                        opt_para = doc.add_paragraph(line.strip())
                        opt_para.paragraph_format.left_indent = Inches(0.5)
            
            # 答案区域（留空或填写）
            if q.get("answer"):
                ans_para = doc.add_paragraph()
                ans_para.add_run(f"答案：{q.get('answer')}").font.size = Pt(11)
            else:
                # 留空
                blank_para = doc.add_paragraph()
                blank_para.add_run("答案：").font.size = Pt(11)
                blank_run = blank_para.add_run("_" * 20)
                blank_run.font.size = Pt(11)
            
            # 添加知识点标签（可选）
            if q.get("knowledge_point"):
                kp_para = doc.add_paragraph()
                kp_run = kp_para.add_run(f"【知识点：{q.get('knowledge_point')}】")
                kp_run.font.size = Pt(9)
                kp_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # 来源信息（可选）
            if q.get("source"):
                src_para = doc.add_paragraph()
                src_run = src_para.add_run(f"来源：{q.get('source')}")
                src_run.font.size = Pt(9)
                src_run.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()  # 题目之间空行
            q_index += 1
    
    # 添加页脚信息
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("— 该试卷由试题库系统生成 —")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 保存文档
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    
    return output_path


def export_to_word_with_answer_sheet(questions: List[Dict], output_path: str, title: str = "高中生物试卷") -> str:
    """
    导出包含答案卷的 Word 文档（题目和答案分开）
    """
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')
    style.font.size = Pt(12)
    
    # ===== 试卷部分 =====
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_run = info_para.add_run("姓名：___________    学号：___________    得分：___________")
    info_run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    questions_by_type = {}
    for q in questions:
        qtype = q.get("question_type", "选择题")
        if qtype not in questions_by_type:
            questions_by_type[qtype] = []
        questions_by_type[qtype].append(q)
    
    q_index = 1
    for qtype, type_questions in questions_by_type.items():
        type_heading = doc.add_heading(qtype, level=1)
        
        for q in type_questions:
            q_para = doc.add_paragraph()
            q_text = f"{q_index}. "
            
            difficulty = q.get("difficulty", 3)
            if difficulty:
                stars = "★" * difficulty + "☆" * (5 - difficulty)
                q_text += f"[{stars}] "
            
            q_text += q.get("title", "")
            q_run = q_para.add_run(q_text)
            q_run.font.size = Pt(12)
            
            if qtype == "选择题" and "\n" in q.get("title", ""):
                lines = q.get("title", "").split("\n")
                for line in lines[1:]:
                    if line.strip():
                        opt_para = doc.add_paragraph(line.strip())
                        opt_para.paragraph_format.left_indent = Inches(0.5)
            
            # 留空答题区域
            if qtype != "选择题":
                blank_para = doc.add_paragraph()
                for _ in range(3):
                    blank_para.add_run("_" * 50 + "\n")
            
            doc.add_paragraph()
            q_index += 1
    
    # 分页
    doc.add_page_break()
    
    # ===== 答案部分 =====
    ans_heading = doc.add_heading("参考答案", level=0)
    ans_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    q_index = 1
    for qtype, type_questions in questions_by_type.items():
        type_heading = doc.add_heading(qtype, level=1)
        
        for q in type_questions:
            ans_para = doc.add_paragraph()
            ans_text = f"{q_index}. {q.get('answer', '（无）')}"
            
            if q.get("knowledge_point"):
                ans_text += f"  【{q.get('knowledge_point')}】"
            
            ans_para.add_run(ans_text).font.size = Pt(11)
            q_index += 1
    
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    
    return output_path


def export_to_excel(questions: List[Dict], output_path: str) -> str:
    """
    导出题目到 Excel 格式（便于编辑）
    """
    import openpyxl
    from openpyxl.styles import Font, Alignment
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "题目列表"
    
    # 表头
    headers = ["序号", "题干", "答案", "难度", "题型", "知识点", "关键词", "来源"]
    ws.append(headers)
    
    # 设置表头样式
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center")
    
    # 写入数据
    for i, q in enumerate(questions, 1):
        difficulty = q.get("difficulty", 3)
        stars = "★" * difficulty + "☆" * (5 - difficulty)
        
        ws.append([
            i,
            q.get("title", ""),
            q.get("answer", ""),
            stars,
            q.get("question_type", "选择题"),
            q.get("knowledge_point", ""),
            q.get("keywords", ""),
            q.get("source", "")
        ])
    
    # 调整列宽
    ws.column_dimensions['A'].width = 6
    ws.column_dimensions['B'].width = 50
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 8
    ws.column_dimensions['E'].width = 10
    ws.column_dimensions['F'].width = 15
    ws.column_dimensions['G'].width = 20
    ws.column_dimensions['H'].width = 15
    
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    wb.save(output_path)
    
    return output_path

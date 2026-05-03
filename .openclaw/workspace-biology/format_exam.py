#!/usr/bin/env python3
"""调整变式训练卷格式：缩小行间距和段间距，排版到6-7页"""

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = '/home/hochoy/.openclaw/media/inbound/å_å¼_è_ç_å_-é_æ_é---1dd64884-48fd-4fc0-9ec4-4fbadfd74339.docx'
OUT = '/home/hochoy/.openclaw/workspace-biology/变式训练卷-排版调整版.docx'

doc = Document(SRC)
ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ===== 页面设置 =====
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

# ===== 全局字体设置 =====
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_spacing(para, before='0', after='0', line='276', lineRule='auto'):
    pPr = para._element.find(f'{{{ns_w}}}pPr')
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    # Remove old spacing elements
    for sp in list(pPr):
        if 'spacing' in sp.tag:
            pPr.remove(sp)
    sp_new = OxmlElement('w:spacing')
    sp_new.set(f'{{{ns_w}}}before', before)
    sp_new.set(f'{{{ns_w}}}after', after)
    sp_new.set(f'{{{ns_w}}}line', line)
    sp_new.set(f'{{{ns_w}}}lineRule', lineRule)
    pPr.append(sp_new)

def is_main_title(text):
    return '变式训练卷' in text and '一、' not in text and '二、' not in text

def is_heading1(text):
    return any(k in text for k in ['一、单项选择题', '二、解答题'])

def is_heading2(text):
    return any(text.startswith(f'{n}．') and len(text) < 15 for n in range(41, 47))

def is_choice_header(text):
    # Q1, Q2 etc. header lines
    import re
    return bool(re.match(r'^\d+．.+（\s+）$', text))

# ===== 处理每个段落 =====
for para in doc.paragraphs:
    text = para.text.strip()
    
    if is_main_title(text):
        # 主标题：段前0，段后6pt，行距1.5
        set_spacing(para, before='0', after='60', line='360')
    elif is_heading1(text):
        # 一级标题：段前10pt，段后4pt，行距1.15
        set_spacing(para, before='100', after='40', line='276')
    elif is_heading2(text):
        # 大题：段前10pt，段后2pt
        set_spacing(para, before='100', after='20', line='276')
    elif text.startswith('（') and '）' in text:
        # 小问：紧排
        set_spacing(para, before='0', after='0', line='260')
    elif text.startswith('A．') or text.startswith('B．') or text.startswith('C．') or text.startswith('D．'):
        # 选项：紧凑
        set_spacing(para, before='0', after='0', line='260')
    elif is_choice_header(text):
        # 题干：紧凑
        set_spacing(para, before='0', after='0', line='276')
    else:
        # 默认：紧凑
        set_spacing(para, before='0', after='0', line='260')

# ===== 处理表格 =====
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_spacing(p, before='0', after='0', line='240')

doc.save(OUT)
print(f'已保存: {OUT}')
